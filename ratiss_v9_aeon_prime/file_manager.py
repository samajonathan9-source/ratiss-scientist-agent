# -*- coding: utf-8 -*-
"""
================================================================================
          GESTIONNAIRE ET CONVERTISSEUR DE FICHIERS — RATISS V9
================================================================================
Propriété Intellectuelle : JohnKing0 & Architecte Jonathan Evina
Version du Système       : RATISS V9 AEON PRIME - INTEGRATED QUANTUM ECOSYSTEM
ID ORCID de l'Auteur     : 0009-0000-4092-5313
Ancrage DOI Académique   : 10.17605/OSF.IO/6JZMB
================================================================================

Ce module offre des outils souverains d'analyse, d'encapsulation, d'archivage, et
de conversion de fichiers hétérogènes (PDF, Word, HTML, XML, ZIP, etc.) en
textes plats structurés (Markdown ou JSON) utilisables par les agents intelligents
et les modèles de langage sans consommer d'appels LLM externes.

---
MOTEURS ET FONCTIONNALITÉS :
1. ANALYSE ET INDEXATION DES FICHIERS :
   - Parcours récursif d'arborescences de workspace de manière sécurisée.
   - Détection de types de fichiers basée sur la signature binaire et les mimetypes.

2. ARCHIVAGE ET ENCAPSULATION :
   - Compression à chaud multi-fichiers au format ZIP ou TAR.GZ.
   - Extraction unifiée d'archives avec décompression isolée (anti Zip-Slip).

3. CONVERTISSEUR HAUTE FIDÉLITÉ (SANS LLM) :
   - PDF : Extraction de texte brut par pypdf / pdfplumber.
   - DOCX : Extraction de paragraphes et tableaux structurés.
   - HTML / XML : Nettoyage d'arbres DOM par BeautifulSoup4 en conservant les titres.
   - JSON / YAML : Reformatage indenté lisible.
   - ZIP : Indexation et arborescence du contenu compressé.

================================================================================
"""

import os
import sys
import zipfile
import tarfile
import mimetypes
import json
import shutil
import re

# ==============================================================================
# CONFIGURATION ET IMPORTATIONS DE SÉCURITÉ POUR LES DÉPENDANCES SCIENTIFIQUES
# ==============================================================================
# Essaye d'importer les librairies d'extraction ; sinon, fallbacks natifs robustes
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


class FileManager:
    """
    Système unifié d'archivage, d'encapsulation et de conversion de documents.
    """
    def __init__(self, workspace_root: str = "."):
        # Chroot léger pour sécuriser l'arborescence des fichiers
        self.workspace_root = os.path.abspath(workspace_root)

    def _resolve_safe_path(self, relative_path: str) -> str:
        """Résout un chemin absolu sécurisé en interdisant la traversée de répertoire."""
        abs_path = os.path.abspath(os.path.join(self.workspace_root, relative_path))
        if not abs_path.startswith(self.workspace_root):
            raise PermissionError(
                f"[SÉCURITÉ] Tentative de traversée de répertoire détectée sur le chemin: {relative_path}"
            )
        return abs_path

    def list_files(self, sub_dir: str = ".") -> list:
        """
        Liste récursivement les fichiers d'un dossier avec leur taille et métadonnées.
        """
        safe_path = self._resolve_safe_path(sub_dir)
        if not os.path.exists(safe_path):
            raise FileNotFoundError(f"Le répertoire spécifié n'existe pas : {sub_dir}")
            
        file_tree = []
        for root, dirs, files in os.walk(safe_path):
            for file in files:
                full_file_path = os.path.join(root, file)
                rel_file_path = os.path.relpath(full_file_path, self.workspace_root)
                stats = os.stat(full_file_path)
                mtype, _ = mimetypes.guess_type(full_file_path)
                
                file_tree.append({
                    "path": rel_file_path,
                    "name": file,
                    "size_bytes": stats.st_size,
                    "last_modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stats.st_mtime)),
                    "mime_type": mtype or "application/octet-stream"
                })
        return file_tree

    def read_file(self, filename: str) -> bytes:
        """Lit le contenu brut (octets) d'un fichier du workspace."""
        safe_path = self._resolve_safe_path(filename)
        with open(safe_path, "rb") as f:
            return f.read()

    def write_file(self, filename: str, content) -> str:
        """Écrit du contenu texte ou binaire dans un fichier du workspace."""
        safe_path = self._resolve_safe_path(filename)
        os.makedirs(os.path.dirname(safe_path), exist_ok=True)
        
        mode = "wb" if isinstance(content, bytes) else "w"
        encoding = None if isinstance(content, bytes) else "utf-8"
        
        with open(safe_path, mode, encoding=encoding) as f:
            f.write(content)
        return f"Fichier écrit avec succès : {filename}"

    def create_directory(self, path: str) -> str:
        """Crée un dossier dans le workspace sécurisé."""
        safe_path = self._resolve_safe_path(path)
        os.makedirs(safe_path, exist_ok=True)
        return f"Répertoire créé avec succès : {path}"

    def delete_file(self, path: str) -> str:
        """Supprime un fichier ou un répertoire récursivement."""
        safe_path = self._resolve_safe_path(path)
        if not os.path.exists(safe_path):
            raise FileNotFoundError(f"Le fichier ou répertoire spécifié n'existe pas : {path}")
            
        if os.path.isdir(safe_path):
            shutil.rmtree(safe_path)
            return f"Répertoire supprimé avec succès : {path}"
        else:
            os.remove(safe_path)
            return f"Fichier supprimé avec succès : {path}"

    # --------------------------------------------------------------------------
    # ARCHIVAGE ET DÉCOMPRESSION
    # --------------------------------------------------------------------------
    def compress_files(self, file_list: list, output_archive: str) -> str:
        """
        Compresse une liste de fichiers dans une archive ZIP ou TAR.GZ.
        """
        safe_archive_path = self._resolve_safe_path(output_archive)
        os.makedirs(os.path.dirname(safe_archive_path), exist_ok=True)
        
        if output_archive.endswith(".zip"):
            with zipfile.ZipFile(safe_archive_path, 'w', zipfile.ZIP_DEFLATED) as zip_f:
                for file in file_list:
                    safe_file_path = self._resolve_safe_path(file)
                    if os.path.exists(safe_file_path):
                        zip_f.write(safe_file_path, os.path.basename(safe_file_path))
            return f"Archive ZIP créée avec succès : {output_archive}"
            
        elif output_archive.endswith(".tar.gz") or output_archive.endswith(".tgz"):
            with tarfile.open(safe_archive_path, 'w:gz') as tar_f:
                for file in file_list:
                    safe_file_path = self._resolve_safe_path(file)
                    if os.path.exists(safe_file_path):
                        tar_f.add(safe_file_path, arcname=os.path.basename(safe_file_path))
            return f"Archive TAR.GZ créée avec succès : {output_archive}"
        else:
            raise ValueError("Format d'archive non supporté. Utilisez .zip ou .tar.gz")

    def extract_archive(self, archive_path: str, target_dir: str = ".") -> str:
        """
        Extrait en toute sécurité le contenu d'un fichier ZIP ou TAR.GZ.
        """
        safe_archive = self._resolve_safe_path(archive_path)
        safe_target = self._resolve_safe_path(target_dir)
        os.makedirs(safe_target, exist_ok=True)
        
        if archive_path.endswith(".zip"):
            with zipfile.ZipFile(safe_archive, 'r') as zip_f:
                # Vérification de sécurité contre les attaques par traversée (Zip-Slip)
                for member in zip_f.infolist():
                    member_path = os.path.abspath(os.path.join(safe_target, member.filename))
                    if not member_path.startswith(safe_target):
                        raise PermissionError(f"[SÉCURITÉ DETECTÉE] Tentative d'attaque Zip-Slip sur : {member.filename}")
                zip_f.extractall(safe_target)
            return f"Archive ZIP extraite avec succès vers {target_dir}"
            
        elif archive_path.endswith(".tar.gz") or archive_path.endswith(".tgz"):
            with tarfile.open(safe_archive, 'r:gz') as tar_f:
                for member in tar_f.getmembers():
                    member_path = os.path.abspath(os.path.join(safe_target, member.name))
                    if not member_path.startswith(safe_target):
                        raise PermissionError(f"[SÉCURITÉ DETECTÉE] Tentative d'attaque Tar-Slip sur : {member.name}")
                tar_f.extractall(safe_target)
            return f"Archive TAR.GZ extraite avec succès vers {target_dir}"
        else:
            raise ValueError("Format d'archive non reconnu pour extraction.")

    # --------------------------------------------------------------------------
    # CONVERTISSEUR UNIFIÉ VERS TEXTE POUR INTELLIGENCE ARTIFICIELLE
    # --------------------------------------------------------------------------
    def convert_to_text(self, filename: str) -> str:
        """
        Extrait et convertit le contenu de n'importe quel fichier complexe
        en texte structuré ultra-lisible (Markdown / JSON).
        """
        safe_path = self._resolve_safe_path(filename)
        if not os.path.exists(safe_path):
            raise FileNotFoundError(f"Fichier introuvable : {filename}")
            
        _, ext = os.path.splitext(filename.lower())
        
        # 1. Traitement des fichiers PDF
        if ext == ".pdf":
            return self._extract_pdf_text(safe_path)
            
        # 2. Traitement des fichiers Word (DOCX)
        elif ext == ".docx":
            return self._extract_docx_text(safe_path)
            
        # 3. Traitement HTML
        elif ext in [".html", ".htm"]:
            return self._extract_html_text(safe_path)
            
        # 4. Traitement JSON
        elif ext == ".json":
            try:
                with open(safe_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return f"```json\n{json.dumps(data, indent=2, ensure_ascii=False)}\n```"
            except Exception as e:
                return f"[ERREUR PARSING JSON] Impossible de lire le JSON : {e}"
                
        # 5. Traitement d'archives ZIP (Analyse d'index)
        elif ext == ".zip":
            try:
                with zipfile.ZipFile(safe_path, 'r') as zip_f:
                    namelist = zip_f.namelist()
                    meta = {
                        "archive_name": os.path.basename(filename),
                        "total_files": len(namelist),
                        "contents": namelist
                    }
                    return f"# INDEX DE L'ARCHIVE COMPRESSÉE : {os.path.basename(filename)}\n" + json.dumps(meta, indent=2, ensure_ascii=False)
            except Exception as e:
                return f"[ERREUR ZIP] Impossible d'indexer l'archive : {e}"
                
        # 6. Traitement par défaut pour les fichiers texte (Python, Markdown, etc.)
        else:
            try:
                with open(safe_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                    return f"```\n{content}\n```"
            except Exception as e:
                return f"[ERREUR FICHIER] Impossible de lire le fichier texte: {e}"

    def _extract_pdf_text(self, path: str) -> str:
        """Méthode interne d'extraction de texte PDF."""
        if not PYPDF_AVAILABLE:
            return "[FALLBACK] Module 'pypdf' non installé. Impossible d'extraire le contenu du fichier PDF de manière structurée."
        try:
            reader = pypdf.PdfReader(path)
            text_pages = []
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text_pages.append(f"--- PAGE {idx + 1} ---\n{page_text}")
            return "\n\n".join(text_pages)
        except Exception as e:
            return f"[ERREUR PDF] Impossible d'analyser le fichier PDF : {e}"

    def _extract_docx_text(self, path: str) -> str:
        """Méthode interne d'extraction de texte Word DOCX."""
        if not DOCX_AVAILABLE:
            return "[FALLBACK] Module 'python-docx' non installé. Impossible d'extraire le document Word."
        try:
            doc = docx.Document(path)
            text_runs = []
            # Extraction des paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_runs.append(paragraph.text)
            # Extraction des tableaux
            for table in doc.tables:
                text_runs.append("\n--- TABLEAU STRUCTURÉ ---")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    text_runs.append(" | ".join(row_data))
            return "\n".join(text_runs)
        except Exception as e:
            return f"[ERREUR DOCX] Impossible d'analyser le document Word : {e}"

    def _extract_html_text(self, path: str) -> str:
        """Méthode interne d'extraction et nettoyage de texte HTML."""
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                html_content = f.read()
            if not BS4_AVAILABLE:
                # Nettoyage par regex basique si bs4 est absent
                cleaned = re.sub(r'<[^>]+>', '', html_content)
                return cleaned
            soup = BeautifulSoup(html_content, "html.parser")
            # Élimine scripts et styles
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator="\n", strip=True)
        except Exception as e:
            return f"[ERREUR HTML] Extraction HTML impossible : {e}"


import time
if __name__ == "__main__":
    fm = FileManager()
    # Création d'un sous-dossier de démo
    fm.create_directory("demo_files")
    fm.write_file("demo_files/test.txt", "RATISS V9 Aeon Prime - Sovereign Physics Kernel")
    print(fm.list_files("demo_files"))
    fm.compress_files(["demo_files/test.txt"], "demo_files/test_archive.zip")
    print(fm.convert_to_text("demo_files/test_archive.zip"))
    fm.delete_file("demo_files")
